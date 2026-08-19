#!/usr/bin/env python3
"""
培训材料 Word 文档导出模块
支持：培训大纲、培训讲义、操作手册、案例分析、测试题、学习笔记模板
输出：专业排版的 .docx 文件
"""

import sys
import json
import argparse
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("需要安装 python-docx：pip install python-docx")
    sys.exit(1)


# ==================== 样式配置 ====================

STYLE_CONFIG = {
    "colors": {
        "primary": RGBColor(0x1A, 0x56, 0xDB),      # 主色 - 蓝色
        "secondary": RGBColor(0x2D, 0x37, 0x48),     # 辅色 - 深灰
        "accent": RGBColor(0xE8, 0x6C, 0x00),        # 强调色 - 橙色
        "success": RGBColor(0x0E, 0x9F, 0x6E),       # 成功色 - 绿色
        "danger": RGBColor(0xE5, 0x3E, 0x3E),        # 警示色 - 红色
        "text": RGBColor(0x1F, 0x29, 0x37),          # 正文色
        "light_text": RGBColor(0x64, 0x74, 0x8B),    # 浅色文字
        "bg_header": RGBColor(0xF0, 0xF4, 0xFF),     # 表头背景
        "white": RGBColor(0xFF, 0xFF, 0xFF),
    },
    "fonts": {
        "cover_title": Pt(28),
        "cover_subtitle": Pt(16),
        "h1": Pt(22),
        "h2": Pt(16),
        "h3": Pt(13),
        "body": Pt(11),
        "small": Pt(9),
        "table_header": Pt(10),
        "table_body": Pt(10),
    },
    "margins": {
        "top": Cm(2.54),
        "bottom": Cm(2.54),
        "left": Cm(3.17),
        "right": Cm(3.17),
    }
}


# ==================== 通用工具函数 ====================

def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_paragraph(doc, text, font_size=None, bold=False, color=None,
                         alignment=None, space_after=None, space_before=None,
                         font_name="微软雅黑"):
    """添加带样式的段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if font_size:
        run.font.size = font_size
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def add_table_row(table, cells_data, is_header=False):
    """添加表格行"""
    row = table.add_row()
    for i, cell_text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(cell_text))
        run.font.size = STYLE_CONFIG["fonts"]["table_body"]
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

        if is_header:
            run.font.bold = True
            run.font.color.rgb = STYLE_CONFIG["colors"]["white"]
            set_cell_shading(cell, "1A56DB")
        else:
            run.font.color.rgb = STYLE_CONFIG["colors"]["text"]
    return row


def add_cover_page(doc, title, subtitle_lines):
    """添加封面页"""
    # 空行撑开
    for _ in range(6):
        doc.add_paragraph()

    # 主标题
    add_styled_paragraph(doc, title,
                         font_size=STYLE_CONFIG["fonts"]["cover_title"],
                         bold=True,
                         color=STYLE_CONFIG["colors"]["primary"],
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(20))

    # 分割线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 30)
    run.font.color.rgb = STYLE_CONFIG["colors"]["accent"]
    run.font.size = Pt(14)

    # 副标题信息
    for line in subtitle_lines:
        add_styled_paragraph(doc, line,
                             font_size=STYLE_CONFIG["fonts"]["cover_subtitle"],
                             color=STYLE_CONFIG["colors"]["secondary"],
                             alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_after=Pt(8))

    # 日期
    add_styled_paragraph(doc, f"编制日期：{datetime.now().strftime('%Y年%m月%d日')}",
                         font_size=STYLE_CONFIG["fonts"]["cover_subtitle"],
                         color=STYLE_CONFIG["colors"]["light_text"],
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(30))

    doc.add_page_break()


def add_toc_page(doc):
    """添加目录页占位"""
    add_styled_paragraph(doc, "目  录",
                         font_size=STYLE_CONFIG["fonts"]["h1"],
                         bold=True,
                         color=STYLE_CONFIG["colors"]["primary"],
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(20))
    add_styled_paragraph(doc, u"(\u76ee\u5f55\u8bf7\u5728Word\u4e2d\u901a\u8fc7 \u5f15\u7528-\u76ee\u5f55-\u63d2\u5165\u76ee\u5f55 \u81ea\u52a8\u751f\u6210)",
                         font_size=STYLE_CONFIG["fonts"]["body"],
                         color=STYLE_CONFIG["colors"]["light_text"],
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def setup_page_style(doc):
    """设置页面样式"""
    for section in doc.sections:
        section.top_margin = STYLE_CONFIG["margins"]["top"]
        section.bottom_margin = STYLE_CONFIG["margins"]["bottom"]
        section.left_margin = STYLE_CONFIG["margins"]["left"]
        section.right_margin = STYLE_CONFIG["margins"]["right"]


# ==================== 培训大纲导出 ====================

def export_outline(data, output_path):
    """导出培训大纲为Word"""
    doc = Document()
    setup_page_style(doc)

    d = data.get("data", data)

    # 封面
    title = d.get("training_name", "培训大纲")
    subtitle_lines = [
        f"培训对象：{d.get('target_audience', '待定')}",
        f"培训时长：{d.get('duration', '待定')}",
    ]
    add_cover_page(doc, title, subtitle_lines)
    add_toc_page(doc)

    # 培训目标
    doc.add_heading("一、培训目标", level=1)
    objectives = d.get("objectives", {})
    for obj_type, label in [("knowledge", "知识目标"), ("skill", "技能目标"), ("attitude", "态度目标")]:
        text = objectives.get(obj_type, "")
        if text:
            add_styled_paragraph(doc, f"【{label}】{text}",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 color=STYLE_CONFIG["colors"]["text"],
                                 space_after=Pt(6))

    # 课程模块
    doc.add_heading("二、课程模块", level=1)
    modules = d.get("modules", [])
    if modules:
        # 模块概览表
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["模块编号", "模块名称", "课时(h)", "教学方法", "考核方式"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.font.bold = True
            run.font.size = STYLE_CONFIG["fonts"]["table_header"]
            run.font.color.rgb = STYLE_CONFIG["colors"]["white"]
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
            set_cell_shading(cell, "1A56DB")

        for mod in modules:
            row = table.add_row()
            cells = [
                mod.get("module_id", ""),
                mod.get("module_name", ""),
                str(mod.get("duration_hours", "")),
                mod.get("teaching_method", ""),
                mod.get("assessment", "")
            ]
            for i, val in enumerate(cells):
                cell = row.cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = STYLE_CONFIG["fonts"]["table_body"]
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

        # 模块详情
        for mod in modules:
            doc.add_heading(f"模块 {mod.get('module_id', '')}：{mod.get('module_name', '')}", level=2)
            points = mod.get("learning_points", [])
            if points:
                add_styled_paragraph(doc, "学习要点：",
                                     font_size=STYLE_CONFIG["fonts"]["h3"],
                                     bold=True,
                                     color=STYLE_CONFIG["colors"]["secondary"],
                                     space_after=Pt(4))
                for pt in points:
                    add_styled_paragraph(doc, f"  • {pt}",
                                         font_size=STYLE_CONFIG["fonts"]["body"],
                                         color=STYLE_CONFIG["colors"]["text"],
                                         space_after=Pt(2))

    # 考核方案
    doc.add_heading("三、考核方案", level=1)
    assessment = d.get("assessment_plan", "")
    if assessment:
        add_styled_paragraph(doc, assessment,
                             font_size=STYLE_CONFIG["fonts"]["body"],
                             color=STYLE_CONFIG["colors"]["text"])
    else:
        add_styled_paragraph(doc, "（考核方案待补充）",
                             font_size=STYLE_CONFIG["fonts"]["body"],
                             color=STYLE_CONFIG["colors"]["light_text"])

    # 教学资源
    doc.add_heading("四、教学资源", level=1)
    add_styled_paragraph(doc, "（教材、设备、软件环境等资源清单待补充）",
                         font_size=STYLE_CONFIG["fonts"]["body"],
                         color=STYLE_CONFIG["colors"]["light_text"])

    doc.save(output_path)
    return output_path


# ==================== 培训讲义导出 ====================

def export_lecture(data, output_path):
    """导出培训讲义为Word"""
    doc = Document()
    setup_page_style(doc)

    d = data.get("data", data)

    # 封面
    title = d.get("module_name", "培训讲义")
    subtitle_lines = [
        f"预计学时：{d.get('duration_hours', '待定')}小时",
        f"前置知识：{d.get('prerequisites', '无')}",
    ]
    add_cover_page(doc, title, subtitle_lines)

    # 学习目标
    doc.add_heading("学习目标", level=1)
    objectives = d.get("learning_objectives", [])
    for obj in objectives:
        add_styled_paragraph(doc, f"✓ {obj}",
                             font_size=STYLE_CONFIG["fonts"]["body"],
                             color=STYLE_CONFIG["colors"]["success"],
                             space_after=Pt(4))

    # 模块概述
    overview = d.get("overview", "")
    if overview:
        doc.add_heading("模块概述", level=1)
        add_styled_paragraph(doc, overview,
                             font_size=STYLE_CONFIG["fonts"]["body"],
                             color=STYLE_CONFIG["colors"]["text"])

    # 核心知识点
    doc.add_heading("核心知识点", level=1)
    concepts = d.get("key_concepts", [])
    for i, concept in enumerate(concepts, 1):
        name = concept.get("name", f"知识点{i}")
        doc.add_heading(f"{i}. {name}", level=2)

        definition = concept.get("definition", "")
        if definition:
            add_styled_paragraph(doc, "【定义】",
                                 font_size=STYLE_CONFIG["fonts"]["h3"],
                                 bold=True,
                                 color=STYLE_CONFIG["colors"]["secondary"])
            add_styled_paragraph(doc, definition,
                                 font_size=STYLE_CONFIG["fonts"]["body"])

        example = concept.get("example", "")
        if example:
            add_styled_paragraph(doc, "【示例】",
                                 font_size=STYLE_CONFIG["fonts"]["h3"],
                                 bold=True,
                                 color=STYLE_CONFIG["colors"]["accent"])
            # 代码/示例框
            p = doc.add_paragraph()
            run = p.add_run(example)
            run.font.size = STYLE_CONFIG["fonts"]["body"]
            run.font.name = "Consolas"
            run.font.color.rgb = STYLE_CONFIG["colors"]["secondary"]

        tips = concept.get("tips", "")
        if tips:
            add_styled_paragraph(doc, f"⚠ 注意：{tips}",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 color=STYLE_CONFIG["colors"]["danger"])

    # 实操步骤
    hands_on = d.get("hands_on_steps", [])
    if hands_on:
        doc.add_heading("实操演练", level=1)
        for step in hands_on:
            step_num = step.get("step_num", "")
            action = step.get("action", "")
            add_styled_paragraph(doc, f"步骤 {step_num}：{action}",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 bold=True,
                                 color=STYLE_CONFIG["colors"]["primary"],
                                 space_after=Pt(2))
            expected = step.get("expected_result", "")
            if expected:
                add_styled_paragraph(doc, f"  ✓ 预期结果：{expected}",
                                     font_size=STYLE_CONFIG["fonts"]["body"],
                                     color=STYLE_CONFIG["colors"]["success"],
                                     space_after=Pt(2))
            trouble = step.get("troubleshooting", "")
            if trouble:
                add_styled_paragraph(doc, f"  ⚠ 常见问题：{trouble}",
                                     font_size=STYLE_CONFIG["fonts"]["body"],
                                     color=STYLE_CONFIG["colors"]["danger"],
                                     space_after=Pt(6))

    # 课堂练习
    exercises = d.get("exercises", [])
    if exercises:
        doc.add_heading("课堂练习", level=1)
        for i, ex in enumerate(exercises, 1):
            ex_type = ex.get("type", "")
            question = ex.get("question", "")
            add_styled_paragraph(doc, f"练习{i}（{ex_type}）",
                                 font_size=STYLE_CONFIG["fonts"]["h3"],
                                 bold=True,
                                 color=STYLE_CONFIG["colors"]["secondary"])
            add_styled_paragraph(doc, question,
                                 font_size=STYLE_CONFIG["fonts"]["body"])

            hint = ex.get("hint", "")
            if hint:
                add_styled_paragraph(doc, f"💡 提示：{hint}",
                                     font_size=STYLE_CONFIG["fonts"]["body"],
                                     color=STYLE_CONFIG["colors"]["light_text"])

    # 要点回顾
    summary = d.get("summary", "")
    if summary:
        doc.add_heading("要点回顾", level=1)
        add_styled_paragraph(doc, summary,
                             font_size=STYLE_CONFIG["fonts"]["body"],
                             color=STYLE_CONFIG["colors"]["text"])

    # 延伸阅读
    readings = d.get("extended_reading", [])
    if readings:
        doc.add_heading("延伸阅读", level=1)
        for r in readings:
            add_styled_paragraph(doc, f"📖 {r}",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 color=STYLE_CONFIG["colors"]["light_text"])

    doc.save(output_path)
    return output_path


# ==================== 操作手册导出 ====================

def export_manual(data, output_path):
    """导出操作手册为Word"""
    doc = Document()
    setup_page_style(doc)

    d = data.get("data", data)

    # 封面
    title = d.get("manual_name", "操作手册")
    subtitle_lines = [
        f"适用系统版本：{d.get('system_version', '待定')}",
        f"目标读者：{d.get('audience', '待定')}",
    ]
    add_cover_page(doc, title, subtitle_lines)
    add_toc_page(doc)

    # 手册用途（支持自定义标题）
    section_titles = d.get("section_titles", {})
    doc.add_heading(section_titles.get("purpose", "一、手册用途"), level=1)
    add_styled_paragraph(doc, d.get("purpose", ""),
                         font_size=STYLE_CONFIG["fonts"]["body"])

    # 使用前提
    doc.add_heading(section_titles.get("prerequisites", "二、使用前提"), level=1)
    prerequisites = d.get("prerequisites", [])
    for pre in prerequisites:
        add_styled_paragraph(doc, f"  • {pre}",
                             font_size=STYLE_CONFIG["fonts"]["body"],
                             space_after=Pt(2))

    # 流程总览
    overview = d.get("process_overview", "")
    if overview:
        doc.add_heading(section_titles.get("process_overview", "三、操作流程总览"), level=1)
        add_styled_paragraph(doc, overview,
                             font_size=STYLE_CONFIG["fonts"]["body"])

    # 操作步骤
    doc.add_heading(section_titles.get("procedures", "四、操作步骤"), level=1)
    procedures = d.get("procedures", [])
    for proc in procedures:
        proc_name = proc.get("proc_name", proc.get("proc_id", "操作"))
        doc.add_heading(f"{proc.get('proc_id', '')} {proc_name}", level=2)

        objective = proc.get("objective", "")
        if objective:
            add_styled_paragraph(doc, f"目的：{objective}",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 color=STYLE_CONFIG["colors"]["secondary"])

        steps = proc.get("steps", [])
        for step in steps:
            step_num = step.get("step_num", "")
            action = step.get("action", "")
            add_styled_paragraph(doc, f"步骤{step_num}：{action}",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 bold=True,
                                 color=STYLE_CONFIG["colors"]["primary"],
                                 space_after=Pt(2))

            screenshot = step.get("screenshot_note", "")
            if screenshot:
                add_styled_paragraph(doc, f"  [截图位置：{screenshot}]",
                                     font_size=STYLE_CONFIG["fonts"]["small"],
                                     color=STYLE_CONFIG["colors"]["light_text"],
                                     space_after=Pt(2))

            expected = step.get("expected", "")
            if expected:
                add_styled_paragraph(doc, f"  ✓ 预期结果：{expected}",
                                     font_size=STYLE_CONFIG["fonts"]["body"],
                                     color=STYLE_CONFIG["colors"]["success"],
                                     space_after=Pt(2))

            tips = step.get("tips", "")
            if tips:
                add_styled_paragraph(doc, f"  ⚠ 注意：{tips}",
                                     font_size=STYLE_CONFIG["fonts"]["body"],
                                     color=STYLE_CONFIG["colors"]["accent"],
                                     space_after=Pt(4))

        verification = proc.get("verification", "")
        if verification:
            add_styled_paragraph(doc, f"验证标准：{verification}",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 bold=True,
                                 color=STYLE_CONFIG["colors"]["success"],
                                 space_before=Pt(8))

    # FAQ
    faq = d.get("faq", [])
    if faq:
        doc.add_heading(section_titles.get("faq", "五、常见问题（FAQ）"), level=1)
        for i, f in enumerate(faq, 1):
            q = f.get("question", "")
            cause = f.get("cause", "")
            solution = f.get("solution", "")

            add_styled_paragraph(doc, f"Q{i}：{q}",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 bold=True,
                                 color=STYLE_CONFIG["colors"]["secondary"])
            if cause:
                add_styled_paragraph(doc, f"  原因：{cause}",
                                     font_size=STYLE_CONFIG["fonts"]["body"])
            add_styled_paragraph(doc, f"  解决：{solution}",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 color=STYLE_CONFIG["colors"]["success"],
                                 space_after=Pt(8))

    # 异常处理
    error_handling = d.get("error_handling", "")
    if error_handling:
        doc.add_heading(section_titles.get("error_handling", "六、异常处理"), level=1)
        add_styled_paragraph(doc, error_handling,
                             font_size=STYLE_CONFIG["fonts"]["body"])

    # 检查清单
    checklist = d.get("checklist", [])
    if checklist:
        doc.add_heading(section_titles.get("checklist", "七、操作检查清单"), level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        for i, h in enumerate(["序号", "检查项", "完成状态"]):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.font.bold = True
            run.font.size = STYLE_CONFIG["fonts"]["table_header"]
            run.font.color.rgb = STYLE_CONFIG["colors"]["white"]
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
            set_cell_shading(cell, "1A56DB")

        for idx, item in enumerate(checklist, 1):
            row = table.add_row()
            for j, val in enumerate([str(idx), item, "☐"]):
                cell = row.cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = STYLE_CONFIG["fonts"]["table_body"]
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    doc.save(output_path)
    return output_path


# ==================== 案例分析导出 ====================

def export_case(data, output_path):
    """导出案例分析为Word"""
    doc = Document()
    setup_page_style(doc)

    d = data.get("data", data)

    # 封面
    title = d.get("case_title", "案例分析")
    case_type = d.get("case_type", "")
    difficulty = d.get("difficulty", "")
    subtitle_lines = [
        f"案例类型：{case_type}",
        f"难度级别：{difficulty}",
        f"业务领域：{d.get('domain', '')}",
    ]
    add_cover_page(doc, title, subtitle_lines)

    # 背景
    doc.add_heading("一、案例背景", level=1)
    add_styled_paragraph(doc, d.get("background", ""),
                         font_size=STYLE_CONFIG["fonts"]["body"])

    # 问题/挑战
    doc.add_heading("二、问题与挑战", level=1)
    challenge = d.get("challenge", {})
    add_styled_paragraph(doc, challenge.get("description", ""),
                         font_size=STYLE_CONFIG["fonts"]["body"])

    key_factors = challenge.get("key_factors", [])
    if key_factors:
        add_styled_paragraph(doc, "关键因素：",
                             font_size=STYLE_CONFIG["fonts"]["h3"],
                             bold=True,
                             color=STYLE_CONFIG["colors"]["secondary"])
        for f in key_factors:
            add_styled_paragraph(doc, f"  • {f}",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 space_after=Pt(2))

    constraints = challenge.get("constraints", "")
    if constraints:
        add_styled_paragraph(doc, f"约束条件：{constraints}",
                             font_size=STYLE_CONFIG["fonts"]["body"],
                             color=STYLE_CONFIG["colors"]["accent"])

    # 决策点
    decision_points = d.get("decision_points", [])
    if decision_points:
        doc.add_heading("三、关键决策点", level=1)
        for dp in decision_points:
            dp_id = dp.get("point_id", "")
            situation = dp.get("situation", "")
            doc.add_heading(f"决策点 {dp_id}", level=2)
            add_styled_paragraph(doc, situation,
                                 font_size=STYLE_CONFIG["fonts"]["body"])

            options = dp.get("options", [])
            if options:
                add_styled_paragraph(doc, "可选方案：",
                                     font_size=STYLE_CONFIG["fonts"]["h3"],
                                     bold=True,
                                     color=STYLE_CONFIG["colors"]["secondary"])
                for opt in options:
                    add_styled_paragraph(doc, f"  ○ {opt}",
                                         font_size=STYLE_CONFIG["fonts"]["body"],
                                         space_after=Pt(2))

            analysis = dp.get("analysis", "")
            if analysis:
                add_styled_paragraph(doc, f"分析：{analysis}",
                                     font_size=STYLE_CONFIG["fonts"]["body"])

            recommended = dp.get("recommended", "")
            if recommended:
                add_styled_paragraph(doc, f"✓ 推荐方案：{recommended}",
                                     font_size=STYLE_CONFIG["fonts"]["body"],
                                     bold=True,
                                     color=STYLE_CONFIG["colors"]["success"])

    # 结果与复盘
    doc.add_heading("四、结果与复盘", level=1)
    outcome = d.get("outcome", {})
    add_styled_paragraph(doc, f"实际结果：{outcome.get('actual_result', '')}",
                         font_size=STYLE_CONFIG["fonts"]["body"])

    lessons = outcome.get("lessons_learned", [])
    if lessons:
        add_styled_paragraph(doc, "经验教训：",
                             font_size=STYLE_CONFIG["fonts"]["h3"],
                             bold=True,
                             color=STYLE_CONFIG["colors"]["danger"])
        for l in lessons:
            add_styled_paragraph(doc, f"  • {l}",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 space_after=Pt(2))

    best = outcome.get("best_practices", [])
    if best:
        add_styled_paragraph(doc, "最佳实践：",
                             font_size=STYLE_CONFIG["fonts"]["h3"],
                             bold=True,
                             color=STYLE_CONFIG["colors"]["success"])
        for b in best:
            add_styled_paragraph(doc, f"  ✓ {b}",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 space_after=Pt(2))

    # 讨论问题
    doc.add_heading("五、课堂讨论", level=1)
    discussion = d.get("discussion", {})
    questions = discussion.get("questions", [])
    for i, q in enumerate(questions, 1):
        add_styled_paragraph(doc, f"{i}. {q}",
                             font_size=STYLE_CONFIG["fonts"]["body"],
                             space_after=Pt(4))

    takeaway = discussion.get("key_takeaway", "")
    if takeaway:
        add_styled_paragraph(doc, f"核心要点：{takeaway}",
                             font_size=STYLE_CONFIG["fonts"]["body"],
                             bold=True,
                             color=STYLE_CONFIG["colors"]["primary"],
                             space_before=Pt(12))

    doc.save(output_path)
    return output_path


# ==================== 测试题导出 ====================

def export_quiz(data, output_path):
    """导出测试题为Word（含试卷和答案两大部分）"""
    doc = Document()
    setup_page_style(doc)

    d = data.get("data", data)

    # ===== 试卷部分 =====
    title = d.get("exam_title", "培训测试")
    subtitle_lines = [
        f"总分：{d.get('total_score', 100)}分",
        f"时间：{d.get('duration_min', 60)}分钟",
        f"及格线：{d.get('passing_score', 60)}分",
    ]
    add_cover_page(doc, title, subtitle_lines)

    # 考生信息
    doc.add_heading("考生信息", level=1)
    info_table = doc.add_table(rows=3, cols=4)
    info_table.style = 'Table Grid'
    info_data = [
        ["姓名", "", "部门", ""],
        ["岗位", "", "工号", ""],
        ["考试日期", "", "得分", ""]
    ]
    for i, row_data in enumerate(info_data):
        for j, val in enumerate(row_data):
            cell = info_table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = STYLE_CONFIG["fonts"]["table_body"]
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
            if j % 2 == 0:
                run.font.bold = True
                set_cell_shading(cell, "F0F4FF")

    doc.add_paragraph()  # 间距

    questions = d.get("questions", {})
    q_num = 0

    # 单选题
    single = questions.get("single_choice", [])
    if single:
        doc.add_heading("一、单选题", level=1)
        for q in single:
            q_num += 1
            add_styled_paragraph(doc, f"{q_num}. {q.get('question', '')}（{q.get('difficulty', '')}）",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 bold=True,
                                 space_after=Pt(2))
            for opt in q.get("options", []):
                add_styled_paragraph(doc, f"    {opt}",
                                     font_size=STYLE_CONFIG["fonts"]["body"],
                                     space_after=Pt(1))
            doc.add_paragraph()

    # 多选题
    multi = questions.get("multiple_choice", [])
    if multi:
        doc.add_heading("二、多选题", level=1)
        for q in multi:
            q_num += 1
            add_styled_paragraph(doc, f"{q_num}. {q.get('question', '')}（{q.get('difficulty', '')}）",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 bold=True,
                                 space_after=Pt(2))
            for opt in q.get("options", []):
                add_styled_paragraph(doc, f"    {opt}",
                                     font_size=STYLE_CONFIG["fonts"]["body"],
                                     space_after=Pt(1))
            doc.add_paragraph()

    # 判断题
    tf = questions.get("true_false", [])
    if tf:
        doc.add_heading("三、判断题", level=1)
        for q in tf:
            q_num += 1
            add_styled_paragraph(doc, f"{q_num}. {q.get('question', '')}（  对  /  错  ）",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 space_after=Pt(6))

    # 简答题
    short = questions.get("short_answer", [])
    if short:
        doc.add_heading("四、简答题", level=1)
        for q in short:
            q_num += 1
            add_styled_paragraph(doc, f"{q_num}. {q.get('question', '')}",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 bold=True,
                                 space_after=Pt(2))
            # 答题空间
            for _ in range(4):
                doc.add_paragraph()

    # 实操题
    practical = questions.get("practical", [])
    if practical:
        doc.add_heading("五、实操题", level=1)
        for q in practical:
            q_num += 1
            add_styled_paragraph(doc, f"{q_num}. {q.get('question', '')}",
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 bold=True,
                                 space_after=Pt(2))
            reqs = q.get("requirements", [])
            for req in reqs:
                add_styled_paragraph(doc, f"    要求：{req}",
                                     font_size=STYLE_CONFIG["fonts"]["body"],
                                     color=STYLE_CONFIG["colors"]["accent"])
            # 答题空间
            for _ in range(6):
                doc.add_paragraph()

    # ===== 答案部分 =====
    doc.add_page_break()
    doc.add_heading("参考答案与解析", level=1)
    add_styled_paragraph(doc, "（本部分仅供考官参考，请勿随试卷发放）",
                         font_size=STYLE_CONFIG["fonts"]["body"],
                         color=STYLE_CONFIG["colors"]["danger"],
                         bold=True,
                         space_after=Pt(16))

    ans_num = 0

    for section_name, section_key in [("单选题", "single_choice"), ("多选题", "multiple_choice")]:
        section_qs = questions.get(section_key, [])
        if section_qs:
            doc.add_heading(f"{section_name}答案", level=2)
            for q in section_qs:
                ans_num += 1
                answer = q.get("answer", "")
                if isinstance(answer, list):
                    answer = ", ".join(answer)
                explanation = q.get("explanation", "")
                kp = q.get("knowledge_point", "")
                add_styled_paragraph(doc, f"{ans_num}. 答案：{answer}",
                                     font_size=STYLE_CONFIG["fonts"]["body"],
                                     bold=True,
                                     color=STYLE_CONFIG["colors"]["primary"])
                if explanation:
                    add_styled_paragraph(doc, f"   解析：{explanation}",
                                         font_size=STYLE_CONFIG["fonts"]["body"],
                                         color=STYLE_CONFIG["colors"]["text"])
                if kp:
                    add_styled_paragraph(doc, f"   知识点：{kp}",
                                         font_size=STYLE_CONFIG["fonts"]["small"],
                                         color=STYLE_CONFIG["colors"]["light_text"])

    for section_name, section_key in [("判断题", "true_false")]:
        section_qs = questions.get(section_key, [])
        if section_qs:
            doc.add_heading(f"{section_name}答案", level=2)
            for q in section_qs:
                ans_num += 1
                answer = "对" if q.get("answer", False) else "错"
                explanation = q.get("explanation", "")
                add_styled_paragraph(doc, f"{ans_num}. 答案：{answer}",
                                     font_size=STYLE_CONFIG["fonts"]["body"],
                                     bold=True,
                                     color=STYLE_CONFIG["colors"]["primary"])
                if explanation:
                    add_styled_paragraph(doc, f"   解析：{explanation}",
                                         font_size=STYLE_CONFIG["fonts"]["body"])

    for section_name, section_key in [("简答题", "short_answer")]:
        section_qs = questions.get(section_key, [])
        if section_qs:
            doc.add_heading(f"{section_name}参考答案", level=2)
            for q in section_qs:
                ans_num += 1
                add_styled_paragraph(doc, f"{ans_num}. 参考答案：",
                                     font_size=STYLE_CONFIG["fonts"]["body"],
                                     bold=True,
                                     color=STYLE_CONFIG["colors"]["primary"])
                add_styled_paragraph(doc, f"   {q.get('answer', '')}",
                                     font_size=STYLE_CONFIG["fonts"]["body"])
                criteria = q.get("scoring_criteria", "")
                if criteria:
                    add_styled_paragraph(doc, f"   评分标准：{criteria}",
                                         font_size=STYLE_CONFIG["fonts"]["body"],
                                         color=STYLE_CONFIG["colors"]["accent"])

    for section_name, section_key in [("实操题", "practical")]:
        section_qs = questions.get(section_key, [])
        if section_qs:
            doc.add_heading(f"{section_name}评分标准", level=2)
            for q in section_qs:
                ans_num += 1
                add_styled_paragraph(doc, f"{ans_num}. 评分标准：",
                                     font_size=STYLE_CONFIG["fonts"]["body"],
                                     bold=True,
                                     color=STYLE_CONFIG["colors"]["primary"])
                add_styled_paragraph(doc, f"   {q.get('scoring_criteria', '')}",
                                     font_size=STYLE_CONFIG["fonts"]["body"])
                expected = q.get("expected_output", "")
                if expected:
                    add_styled_paragraph(doc, f"   预期输出：{expected}",
                                         font_size=STYLE_CONFIG["fonts"]["body"],
                                         color=STYLE_CONFIG["colors"]["success"])

    doc.save(output_path)
    return output_path


# ==================== 学习笔记模板导出 ====================

def export_notes_template(data, output_path):
    """导出学习笔记模板为Word"""
    doc = Document()
    setup_page_style(doc)

    d = data.get("data", data)
    template_type = d.get("template_type", "cornell")
    course_name = d.get("course_name", "")

    if template_type == "cornell":
        title = f"康奈尔笔记模板"
        if course_name:
            title += f" — {course_name}"
        add_cover_page(doc, title, ["基于康奈尔笔记法", "左侧线索 | 右侧笔记 | 底部总结"])

        doc.add_heading("使用说明", level=1)
        instructions = [
            "1. 记录（右侧栏）：上课时记录关键内容和要点",
            "2. 线索（左侧栏）：课后提炼关键词和问题",
            "3. 总结（底部）：用自己的话概括本页核心内容",
            "4. 复习：遮住右侧，看左侧线索回忆内容"
        ]
        for inst in instructions:
            add_styled_paragraph(doc, inst,
                                 font_size=STYLE_CONFIG["fonts"]["body"],
                                 space_after=Pt(4))

        # 生成笔记模板页
        for page_num in range(1, 6):
            doc.add_heading(f"笔记页 {page_num}", level=2)

            # 主题行
            topic_table = doc.add_table(rows=1, cols=2)
            topic_table.style = 'Table Grid'
            cell_left = topic_table.rows[0].cells[0]
            cell_right = topic_table.rows[0].cells[1]
            cell_left.text = ""
            cell_right.text = ""
            p1 = cell_left.paragraphs[0]
            run1 = p1.add_run("主题：")
            run1.font.size = STYLE_CONFIG["fonts"]["body"]
            run1.font.bold = True
            run1.font.name = "微软雅黑"
            run1._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
            set_cell_shading(cell_left, "F0F4FF")

            p2 = cell_right.paragraphs[0]
            run2 = p2.add_run("日期：          讲师：")
            run2.font.size = STYLE_CONFIG["fonts"]["body"]
            run2.font.name = "微软雅黑"
            run2._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
            set_cell_shading(cell_right, "F0F4FF")

            doc.add_paragraph()

            # 主体：线索 | 笔记
            main_table = doc.add_table(rows=8, cols=2)
            main_table.style = 'Table Grid'

            # 设置列宽比例（左侧1/3，右侧2/3）
            for row in main_table.rows:
                row.cells[0].width = Cm(5.5)
                row.cells[1].width = Cm(11)

            # 表头
            for i, h in enumerate(["线索/关键词", "笔记内容"]):
                cell = main_table.rows[0].cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(h)
                run.font.bold = True
                run.font.size = STYLE_CONFIG["fonts"]["table_header"]
                run.font.color.rgb = STYLE_CONFIG["colors"]["white"]
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
                set_cell_shading(cell, "1A56DB")

            # 底部总结
            doc.add_paragraph()
            summary_table = doc.add_table(rows=2, cols=1)
            summary_table.style = 'Table Grid'
            header_cell = summary_table.rows[0].cells[0]
            header_cell.text = ""
            p = header_cell.paragraphs[0]
            run = p.add_run("本页总结")
            run.font.bold = True
            run.font.size = STYLE_CONFIG["fonts"]["table_header"]
            run.font.color.rgb = STYLE_CONFIG["colors"]["white"]
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
            set_cell_shading(header_cell, "E86C00")

            doc.add_page_break()

    elif template_type == "mindmap_style":
        title = f"思维导图式笔记模板"
        if course_name:
            title += f" — {course_name}"
        add_cover_page(doc, title, ["以核心主题为中心，向外辐射分支", "适合概念梳理和知识关联"])

        for page_num in range(1, 6):
            doc.add_heading(f"笔记页 {page_num}", level=2)
            add_styled_paragraph(doc, "核心主题：____________________",
                                 font_size=STYLE_CONFIG["fonts"]["h3"],
                                 bold=True,
                                 color=STYLE_CONFIG["colors"]["primary"],
                                 space_after=Pt(12))

            branches = ["分支1", "分支2", "分支3", "分支4"]
            for branch in branches:
                add_styled_paragraph(doc, f"  {branch}：____________________",
                                     font_size=STYLE_CONFIG["fonts"]["body"],
                                     color=STYLE_CONFIG["colors"]["secondary"],
                                     space_after=Pt(4))
                add_styled_paragraph(doc, "    • 详细内容：",
                                     font_size=STYLE_CONFIG["fonts"]["body"],
                                     space_after=Pt(2))
                add_styled_paragraph(doc, "    • 关联知识：",
                                     font_size=STYLE_CONFIG["fonts"]["body"],
                                     space_after=Pt(8))

            add_styled_paragraph(doc, "关联发现：",
                                 font_size=STYLE_CONFIG["fonts"]["h3"],
                                 bold=True,
                                 color=STYLE_CONFIG["colors"]["accent"])
            add_styled_paragraph(doc, "________________________________",
                                 font_size=STYLE_CONFIG["fonts"]["body"])
            doc.add_page_break()

    elif template_type == "practical_log":
        title = f"实操记录模板"
        if course_name:
            title += f" — {course_name}"
        add_cover_page(doc, title, ["记录实操步骤、问题、解决方案", "适合实验课和实操课"])

        for page_num in range(1, 6):
            doc.add_heading(f"实操记录 {page_num}", level=2)

            sections = [
                ("实操项目", "项目名称："),
                ("操作步骤", "1. \n2. \n3. "),
                ("遇到的问题", ""),
                ("解决方案", ""),
                ("关键收获", ""),
                ("下一步计划", ""),
            ]
            for sec_title, sec_content in sections:
                add_styled_paragraph(doc, f"【{sec_title}】",
                                     font_size=STYLE_CONFIG["fonts"]["h3"],
                                     bold=True,
                                     color=STYLE_CONFIG["colors"]["secondary"],
                                     space_after=Pt(4))
                if sec_content:
                    add_styled_paragraph(doc, sec_content,
                                         font_size=STYLE_CONFIG["fonts"]["body"],
                                         space_after=Pt(8))
                else:
                    doc.add_paragraph()
            doc.add_page_break()

    doc.save(output_path)
    return output_path


# ==================== 路由与主入口 ====================

EXPORTERS = {
    "outline": export_outline,
    "lecture": export_lecture,
    "manual": export_manual,
    "case": export_case,
    "quiz": export_quiz,
    "notes": export_notes_template,
}


def export_to_docx(material_type, data_or_file, output_path):
    """
    通用导出接口
    material_type: outline/lecture/manual/case/quiz/notes
    data_or_file: JSON数据字典 或 JSON文件路径
    output_path: 输出.docx文件路径
    """
    if material_type not in EXPORTERS:
        raise ValueError(f"不支持的材料类型: {material_type}，支持: {', '.join(EXPORTERS.keys())}")

    # 加载数据
    if isinstance(data_or_file, str) or isinstance(data_or_file, Path):
        with open(data_or_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
    else:
        data = data_or_file

    exporter = EXPORTERS[material_type]
    result = exporter(data, str(output_path))
    return result


def main():
    parser = argparse.ArgumentParser(description="培训材料Word文档导出工具")
    parser.add_argument("--type", "-t", required=True,
                        help="材料类型：outline/lecture/manual/case/quiz/notes")
    parser.add_argument("--input", "-i", required=True,
                        help="输入JSON数据文件路径")
    parser.add_argument("--output", "-o", required=True,
                        help="输出.docx文件路径")

    args = parser.parse_args()

    try:
        result = export_to_docx(args.type, args.input, args.output)
        print(f"导出成功: {result}")
    except Exception as e:
        print(f"导出失败: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
